import cgi
import os
import tempfile
from pathlib import Path
from typing import BinaryIO
from urllib.parse import unquote

import docx
import pandas as pd
import requests
from PyPDF2 import PdfReader
from spire.doc import Document as SpireDocument


def load_excel_file(file_path: str) -> pd.DataFrame:
    """Load Excel file into a pandas DataFrame"""
    try:
        return pd.read_excel(file_path)
    except Exception as e:
        print(f"Failed to load Excel file: {e}")
        return None


def download_file(url: str, output_folder: str) -> None:
    """Download a file from a URL and save it to the output folder"""
    try:
        response = requests.get(url, stream=True, allow_redirects=True)
        response.raise_for_status()

        file_name = None
        if "Content-Disposition" in response.headers:
            disposition = response.headers["Content-Disposition"]
            _, params = cgi.parse_header(disposition)
            file_name = params.get("filename*")
            if file_name:
                file_name = file_name.split("''")[-1]
                file_name = unquote(file_name)

        if file_name is None:
            file_name = "downloaded_file"  # default file name

        file_path = Path(output_folder) / file_name
        with open(file_path, "wb") as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        print(f"Downloaded: {file_name}")
    except requests.RequestException as e:
        print(f"Failed to download {url}: {e}")
    except Exception as e:
        print(f"An error occurred: {e}")


def download_files_from_excel(
    file_path: str, url_column_name: str, output_folder: str
) -> None:
    """Read Excel file and download files from URLs"""
    df = load_excel_file(file_path)
    if df is None:
        return

    # Ensure the output folder exists
    Path(output_folder).mkdir(parents=True, exist_ok=True)

    for index, row in df.iterrows():
        url = row[url_column_name]
        download_file(url, output_folder)


def parse_files_to_txt(input_folder: str, output_folder: str) -> None:
    """Parse files from input folder to txts in output folder"""
    # Ensure the output folder exists
    Path(output_folder).mkdir(parents=True, exist_ok=True)

    for file_path in Path(input_folder).glob("*"):
        if file_path.is_file():
            file_name = file_path.name
            file_extension = file_path.suffix.lower()

            if file_extension == ".doc":
                with open(file_path, "rb") as f:
                    text = doc_reader(f)
            elif file_extension == ".docx":
                document = docx.Document(file_path)
                text = "\n".join([paragraph.text for paragraph in document.paragraphs])
            elif file_extension == ".pdf":
                with open(file_path, "rb") as f:
                    pdf_reader = PdfReader(f)
                    text = "\n".join([page.extract_text() for page in pdf_reader.pages])
            else:
                continue

            output_file_path = Path(output_folder) / f"{file_name}.txt"
            with open(output_file_path, "w") as f:
                f.write(text)

            print(f"Parsed {file_name} to txt")


def doc_reader(doc_bytes: BinaryIO) -> str:
    # your existing function
    with tempfile.NamedTemporaryFile(delete=True, suffix=".doc") as tmp:
        tmp.write(doc_bytes.read())
        tmp.seek(0)

        document = SpireDocument()
        document.LoadFromFile(tmp.name)

        text = document.GetText()
        return text[71:]


if __name__ == "__main__":
    print("Current working directory:", os.getcwd())

    file_path = Path("../documents/data.xlsx").resolve()
    url_column_name = "url"
    output_folder = Path("../documents/docs").resolve()
    download_files_from_excel(str(file_path), url_column_name, str(output_folder))

    input_folder = Path("../documents/docs").resolve()
    output_folder = Path("../documents/txts").resolve()
    parse_files_to_txt(str(input_folder), str(output_folder))
