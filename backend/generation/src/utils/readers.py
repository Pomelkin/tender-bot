import tempfile
from io import BytesIO
from typing import BinaryIO

import docx
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
from spire.doc import Document as SpireDocument


class DocumentReader:
    async def _read_pdf(self, bytesio_obj: BytesIO) -> str:
        """
        Reads a PDF from a BytesIO object and returns its text content as a string.
        """
        pdf_reader = PdfReader(bytesio_obj)
        text = ""
        for page in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page].extract_text()
        return text

    async def _read_docx(self, bytesio_obj: BytesIO) -> str:
        """
        Reads a DOCX file from a BytesIO object and returns its text content as a string.
        """
        doc = docx.Document(bytesio_obj)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    async def _doc_reader(self, doc_bytes: BinaryIO) -> str:
        with tempfile.NamedTemporaryFile(delete=True, suffix=".doc") as tmp:
            tmp.write(doc_bytes.read())
            tmp.seek(0)

            document = SpireDocument()
            document.LoadFromFile(tmp.name)

            text = document.GetText()
            return text[71:]

    async def _read_html(self, bytesio_obj: BytesIO) -> str:
        """
        Reads an HTML file from a BytesIO object and returns its text content as a string.
        """
        html_str = bytesio_obj.getvalue().decode("utf-8")
        soup = BeautifulSoup(html_str, "html.parser")
        text = soup.get_text()
        return text

    async def read(self, file_type: str, bytesio_obj: BytesIO) -> str:
        """
        Reads a document from a BytesIO object and returns its text content as a string.
        """
        file_type = file_type.lower()
        if file_type == "pdf":
            return await self._read_pdf(bytesio_obj)
        elif file_type == "docx":
            return await self._read_docx(bytesio_obj)
        elif file_type == "doc":
            return await self._doc_reader(bytesio_obj)
        elif file_type == "html":
            return await self._read_html(bytesio_obj)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")


if __name__ == "__main__":
    # Example usage:
    with open("example.pdf", "rb") as f:
        pdf_bytes = BytesIO(f.read())
